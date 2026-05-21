"""Generate docs/paper_notes.docx — technical paper companion document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement

REPO = Path(__file__).resolve().parents[1]
OUT = REPO / "docs" / "paper_notes.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def note(doc: Document, text: str) -> None:
    """Italic grey note paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def fill(doc: Document, text: str) -> None:
    """[FILL AFTER RUNS] placeholder paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def bullet(doc: Document, text: str, level: int = 0) -> None:
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.4)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
        cell.paragraphs[0].runs[0].bold = True
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()  # spacing after table


def checklist(doc: Document, items: list[tuple[bool, str]]) -> None:
    for done, text in items:
        mark = "☑" if done else "☐"
        doc.add_paragraph(f"{mark}  {text}", style="List Bullet")


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

def build(doc: Document) -> None:
    doc.add_heading("Technical Paper Notes — Roborto", 0)
    note(doc, (
        "Companion document for writing the arXiv technical reports. "
        "All numbers are taken directly from the implemented code and configs — "
        "do not edit them by hand. "
        "Sections marked in red must be filled in after W&B results arrive."
    ))
    doc.add_paragraph()

    # ------------------------------------------------------------------
    h1(doc, "1. Research Framing")
    # ------------------------------------------------------------------

    h2(doc, "1.1 Central Question (Spine A)")
    para(doc, (
        "What makes an embodied agent learn fast and generalise, "
        "and does the answer transfer between manipulation and navigation?"
    ), italic=True)
    doc.add_paragraph()

    para(doc, "Two levers under study:")
    table(doc,
          ["Platform", "Primary lever", "Task family"],
          [
              ["Gymnasium-Robotics (MuJoCo)", "Goal relabeling strategy + exploration",
               "Sparse-reward goal-conditioned manipulation"],
              ["AI Habitat", "Frozen visual encoder quality",
               "Visual navigation (PointNav / ImageNav)"],
          ])

    para(doc, "Cross-cutting synthesis question:", bold=True)
    para(doc, (
        "Is 'good representation / good exploration' a domain-specific "
        "or domain-general property?"
    ), italic=True)

    h2(doc, "1.2 Claim This Paper Can Support (Phase 1 Only)")
    para(doc, (
        "A correctly-configured SAC+HER agent matches published success rates "
        "on sparse-reward Fetch tasks, establishing a trusted baseline from which "
        "the relabeling-strategy study departs. "
        "This is a reproducibility claim + infrastructure claim — state it honestly."
    ))

    h2(doc, "1.3 One-Sentence Positioning")
    para(doc, (
        "We reproduce the SAC+HER baseline on FetchPush, FetchPickAndPlace, and FetchSlide "
        "with proper multi-seed statistics, providing an auditable starting point for the "
        "relabeling-strategy ablation that follows."
    ), italic=True)

    # ------------------------------------------------------------------
    h1(doc, "2. Environment")
    # ------------------------------------------------------------------

    h2(doc, "2.1 FetchPush-v4 (Primary Baseline Task)")
    table(doc,
          ["Property", "Value", "Source"],
          [
              ["Library", "gymnasium-robotics 1.2.3", "environment-gymrobotics.yml"],
              ["Gymnasium ID", "FetchPush-v4", "configs/env/fetch_push.yaml"],
              ["Physics engine", "MuJoCo 3.8.1", "environment-gymrobotics.yml"],
              ["Observation — proprioception", "25-dimensional float64", "verified via env.observation_space"],
              ["Observation — achieved_goal", "3-dimensional (puck XYZ)", "verified via env.observation_space"],
              ["Observation — desired_goal", "3-dimensional (target XYZ)", "verified via env.observation_space"],
              ["Observation total", "31-dimensional (25 + 3 + 3)", "—"],
              ["Action space", "Box(−1, 1, shape=(4,), float32) — 3D Δ-EE + gripper", "verified via env.action_space"],
              ["Reward", "Sparse: 0 if ‖g_a − g_d‖₂ ≤ 0.05 m, else −1", "reward_type: sparse"],
              ["Episode horizon", "50 steps", "max_episode_steps: 50"],
              ["Task", "Push a puck to a target position on a table", "—"],
              ["Goal space", "3D target puck position", "Gymnasium-Robotics docs"],
          ])

    para(doc, "Also planned (same config family, once FetchPush baseline is confirmed):", bold=True)
    bullet(doc, "FetchReach-v4 — sanity check (easiest task, ~100% expected)")
    bullet(doc, "FetchPickAndPlace-v4 — harder (lift + place)")
    bullet(doc, "FetchSlide-v4 — hardest (no direct contact possible)")

    h2(doc, "2.2 Goal-Conditioned RL Formulation")
    para(doc, (
        "Fetch tasks implement the GoalEnv interface (Plappert et al., 2018). "
        "Each observation is a dict with three components:"
    ))
    bullet(doc, "observation s: robot proprioception (joint angles, velocities, object position/velocity)")
    bullet(doc, "achieved_goal g_a: current object position")
    bullet(doc, "desired_goal g_d: sampled per episode target position")
    para(doc, (
        "Success criterion: ‖g_a − g_d‖₂ ≤ δ with δ = 0.05 m.  "
        "Sparse reward: r(g_a, g_d) = −𝟙{‖g_a − g_d‖₂ > δ}, i.e. 0 on success, −1 otherwise."
    ))

    # ------------------------------------------------------------------
    h1(doc, "3. Algorithm")
    # ------------------------------------------------------------------

    h2(doc, "3.1 Soft Actor-Critic (SAC)")
    para(doc, (
        "SAC (Haarnoja et al., 2018) is an off-policy maximum-entropy actor-critic algorithm. "
        "It maximises J(π) = E[∑ γᵗ (r_t + α H(π(·|s_t)))] where α is the entropy temperature "
        "(auto-tuned). Compared to DDPG it is more stable and sample-efficient on continuous "
        "control tasks with sparse rewards."
    ))
    para(doc, (
        "Implementation: Stable-Baselines3 2.8.0, SAC class with MultiInputPolicy. "
        "MultiInputPolicy uses a CombinedExtractor feature extractor "
        "(flat MLP per dict key, outputs concatenated) feeding into a 2-layer MLP "
        "policy and critic."
    ))
    para(doc, "Default network architecture (SB3 2.8.0, net_arch=None):", bold=True)
    para(doc, "Feature extractor output → 256 hidden units → 256 hidden units → action / Q-value")

    h2(doc, "3.2 Hindsight Experience Replay (HER)")
    para(doc, (
        "HER (Andrychowicz et al., 2017) addresses sparse rewards in goal-conditioned "
        "environments by relabeling past transitions with goals that were actually achieved. "
        "For a transition (s_t, a_t, s_{t+1}, g_d) that yielded reward −1, HER substitutes "
        "g_d with g_a(s_{t+k}) for some k, recomputes the reward (now 0 = success), and adds "
        "this synthetic transition to the replay buffer. This provides a dense learning signal "
        "even when the policy almost never reaches the true desired goal."
    ))
    para(doc, "Relabeling strategies:", bold=True)
    table(doc,
          ["Strategy", "Goal selection", "SB3 name"],
          [
              ["future", "Uniform from {t+1, …, T} — future states in same episode", "future"],
              ["final", "Last state of the episode", "final"],
              ["episode", "Uniform from {0, …, T} — any state in same episode", "episode"],
              ["random", "Uniform from the entire replay buffer", "random"],
          ])
    para(doc, (
        "Baseline uses future with k=4 (4 relabeled transitions per real transition), "
        "consistent with Andrychowicz et al. The relabeling-strategy comparison (Phase 1 "
        "contribution) will sweep all four strategies."
    ))
    para(doc, (
        "In-repo implementation: roborto/buffers/her.py — HERRelabeler class implements "
        "all four strategies with vectorised reward recomputation. The baseline uses SB3's "
        "HerReplayBuffer; the contribution experiment swaps in HERRelabeler for "
        "strategy-level control."
    ))

    h2(doc, "3.3 Exact Hyperparameters — exp001 Baseline")
    table(doc,
          ["Parameter", "Value", "Config key"],
          [
              ["Discount factor γ", "0.98", "algo.gamma"],
              ["Learning rate", "1 × 10⁻³", "algo.learning_rate"],
              ["Batch size", "256", "algo.batch_size"],
              ["Replay buffer size", "1 × 10⁶", "algo.buffer_size"],
              ["Soft update τ", "0.005", "algo.tau"],
              ["Learning starts", "1,000 steps", "algo.learning_starts"],
              ["HER strategy", "future", "algo.her.strategy"],
              ["HER relabeling ratio k", "4", "algo.her.n_relabel"],
              ["Policy / critic network", "[256, 256] MLP", "SB3 default (net_arch=None)"],
              ["Entropy temperature α", "Auto-tuned", "SB3 default"],
              ["Actor / critic optimiser", "Adam", "SB3 default"],
              ["Total training steps", "500,000", "trainer.total_steps"],
              ["Number of seeds", "8", "trainer.num_seeds"],
              ["Eval interval", "25,000 steps", "trainer.eval_interval"],
              ["Eval episodes per checkpoint", "100", "trainer.eval_episodes"],
          ])

    para(doc, "Reproduce with one command:", bold=True)
    code(doc, "make train EXP=exp001_fetch_sac_her GPU=GPU-5ac08ef4-9737-c367-bf80-200f205f6014")
    code(doc, "make run-all GPU=GPU-5ac08ef4-9737-c367-bf80-200f205f6014 SEEDS=0,1,2,3,4,5,6,7")

    # ------------------------------------------------------------------
    h1(doc, "4. Evaluation Protocol")
    # ------------------------------------------------------------------

    h2(doc, "4.1 Metrics")
    para(doc, "Primary metric: success rate", bold=True)
    para(doc, (
        "Fraction of evaluation episodes in which the agent achieves the goal "
        "(‖g_a − g_d‖₂ ≤ 0.05) before episode termination. "
        "Implemented in roborto/evaluation/metrics.py."
    ))
    para(doc, "Secondary metric: mean episodic return", bold=True)
    para(doc, (
        "Mean undiscounted sum of rewards over evaluation episodes "
        "(range [−50, 0] for FetchPush with horizon 50 and sparse reward −1)."
    ))
    para(doc, "Sample efficiency:", bold=True)
    para(doc, "Steps-to-X% success — read off learning curves.")

    h2(doc, "4.2 Evaluation Procedure")
    bullet(doc, "Deterministic policy (model.predict(obs, deterministic=True))")
    bullet(doc, "100 fresh episodes per evaluation checkpoint")
    bullet(doc, "Evaluation seed = training seed + 1000 (separate instance, never shared)")
    bullet(doc, "Evaluation happens every 25,000 training steps → 20 data points over a 500k run")

    h2(doc, "4.3 Multi-Seed Aggregation (rliable)")
    para(doc, (
        "Raw per-seed learning curves are logged to W&B. For the paper:"
    ))
    numbered(doc, "Extract all 8 seeds' learning curves at each eval checkpoint.")
    numbered(doc, (
        "Report IQM (Interquartile Mean) ± stratified bootstrap 95% CI using rliable "
        "(Agarwal et al., 2021). IQM is preferred over mean for RL — it is robust to "
        "outlier seeds."
    ))
    numbered(doc, "Plot as a shaded learning curve (IQM ± CI vs training steps).")
    numbered(doc, "Final performance = IQM of success rate at the last checkpoint (step 500,000).")
    para(doc, "Figures are generated by scripts/make_figures.py — never hand-made.")

    h2(doc, "4.4 Target Numbers (Baselines to Match)")
    fill(doc, "Fill in from Andrychowicz et al. 2017 before finalising the paper.")
    table(doc,
          ["Task", "Algorithm", "Target success rate", "Source"],
          [
              ["FetchPush", "SAC+HER (future, k=4)", "[FILL] ~80–90%", "Andrychowicz et al. 2017"],
              ["FetchReach", "SAC+HER", "[FILL] ~100%", "—"],
              ["FetchPickAndPlace", "SAC+HER", "[FILL] ~50–70%", "—"],
              ["FetchSlide", "SAC+HER", "[FILL] ~30–60%", "—"],
          ])
    para(doc, (
        "Success criterion: final success rate matches target within bootstrap CI "
        "on at least 2 tasks."
    ))

    # ------------------------------------------------------------------
    h1(doc, "5. Infrastructure")
    # ------------------------------------------------------------------

    h2(doc, "5.1 Repository Structure")
    code(doc, "roborto/")
    code(doc, "  roborto/              # installable package (pip install -e .[gym])")
    code(doc, "    buffers/her.py      # HERRelabeler — 4 strategies, fully tested")
    code(doc, "    buffers/replay_buffer.py  # GoalReplayBuffer — circular, numpy")
    code(doc, "    envs/gym_robotics/make.py # make_env(cfg, seed)")
    code(doc, "    algos/sac.py        # build_sac(cfg, env, device) — SB3 SAC + HER")
    code(doc, "    training/trainer.py # SACHERTrainer — smoke, resume, eval, ckpt")
    code(doc, "    evaluation/metrics.py     # success_rate(), spl()")
    code(doc, "    utils/              # seeding, device, logging, provenance")
    code(doc, "  configs/              # Hydra config tree — ALL hyperparameters")
    code(doc, "  scripts/train.py      # entry point (Hydra @main)")
    code(doc, "  scripts/launch_seeds.py    # parallel/pooled multi-seed launcher")
    code(doc, "  scripts/run_all_exps.py    # loop all manipulation experiments")
    code(doc, "  tests/               # 24 unit + smoke tests, all green")

    h2(doc, "5.2 Reproducibility Guarantees")
    para(doc, "Every training run automatically records (W&B + provenance.json):")
    bullet(doc, "Git SHA of the commit that produced it")
    bullet(doc, "Full resolved Hydra config (all hyperparameters, no hidden defaults)")
    bullet(doc, "Exact package versions of every installed package (importlib.metadata)")
    bullet(doc, "Global seed set for Python random, NumPy, and PyTorch")

    h2(doc, "5.3 Seeding")
    para(doc, "roborto.utils.seeding.set_global_seed(seed) sets:")
    bullet(doc, "os.environ['PYTHONHASHSEED']")
    bullet(doc, "random.seed")
    bullet(doc, "numpy.random.seed")
    bullet(doc, "torch.manual_seed + torch.cuda.manual_seed_all")
    bullet(doc, "cudnn.deterministic=True, cudnn.benchmark=False")
    para(doc, (
        "Training env seeded via env.reset(seed=cfg.seed). "
        "Evaluation env seeded via env.reset(seed=cfg.seed + 1000) "
        "(separate instance, never shared with training)."
    ))

    h2(doc, "5.4 Checkpointing and Resumption")
    para(doc, "Every eval cycle writes:")
    bullet(doc, "<run_dir>/checkpoint/model.zip — SB3 model (policy + replay buffer)")
    bullet(doc, "<run_dir>/checkpoint/step.json — number of environment steps completed")
    para(doc, (
        "On relaunch with trainer.resume=true (the default), the trainer detects the "
        "checkpoint and resumes from the recorded step. Critical on a shared, "
        "unscheduled GPU box."
    ))

    h2(doc, "5.5 Compute")
    table(doc,
          ["Component", "Hardware", "Notes"],
          [
              ["Manipulation training", "Full NVIDIA A100-SXM4-80GB",
               "MIG slices cause CUDA init errors with SB3 — use full cards only"],
              ["GPU selection", "CUDA_VISIBLE_DEVICES=<UUID>",
               "Set by Makefile; UUIDs stable across reboots"],
              ["Wall time / seed", "~6 hrs at 500k steps",
               "~12 hrs/1M steps empirically measured on A100"],
              ["Total for 8 seeds (sequential)", "~48 hrs", "make run-all on 1 GPU"],
              ["VRAM", "< 2 GB", "SAC on FetchPush — tiny policy network"],
          ])
    para(doc, (
        "Verified GPU: GPU-5ac08ef4-9737-c367-bf80-200f205f6014 (NVIDIA A100-SXM4-80GB). "
        "Smoke run confirmed end-to-end on 2026-05-21."
    ))

    # ------------------------------------------------------------------
    h1(doc, "6. Logging and W&B")
    # ------------------------------------------------------------------

    para(doc, "Project: https://wandb.ai/sohampahari-research-valeo/roborto", bold=True)
    para(doc, "Metrics logged at every eval checkpoint (every 25k steps):")
    bullet(doc, "eval/success_rate — primary metric")
    bullet(doc, "eval/mean_return — secondary metric")
    para(doc, "Run naming: <experiment_name>/seed<N>  e.g. exp001_sac_her_fetch_push/seed3")
    para(doc, (
        "To view all seeds overlaid: W&B → project → Runs → Group by experiment_name."
    ))

    # ------------------------------------------------------------------
    h1(doc, "7. Paper Section Drafts")
    # ------------------------------------------------------------------

    h2(doc, "7.1 Abstract Skeleton")
    para(doc, (
        "We present Roborto, a reproducible research framework for embodied reinforcement "
        "learning across manipulation (Gymnasium-Robotics/MuJoCo) and navigation (AI Habitat). "
        "We reproduce the SAC+HER baseline on sparse-reward Fetch tasks with proper multi-seed "
        "statistics (N=8 seeds, IQM ± 95% CI), achieving "
    ))
    fill(doc, "[FILL: X%] success on FetchPush and [FILL: Y%] on FetchPickAndPlace")
    para(doc, "after 500k environment steps — within the confidence interval of")
    fill(doc, "[FILL: target from Andrychowicz et al. 2017]. We then study [...].")

    h2(doc, "7.2 Methods — Environments Paragraph")
    para(doc, (
        "We evaluate on the Fetch manipulation suite from Gymnasium-Robotics "
        "(Plappert et al., 2018), implemented with MuJoCo 3.8.1. Each task is a "
        "sparse-reward goal-conditioned environment (GoalEnv): the agent receives "
        "observation s ∈ ℝ²⁵, achieved goal g_a ∈ ℝ³, and desired goal g_d ∈ ℝ³, "
        "with binary reward r = −𝟙{‖g_a − g_d‖₂ > 0.05}. "
        "The action space is continuous in ℝ⁴ (3D end-effector displacement + gripper width). "
        "Episode horizon is 50 steps. We report results on FetchPush (primary), FetchReach "
        "(sanity), FetchPickAndPlace, and FetchSlide."
    ))

    h2(doc, "7.3 Methods — Algorithm Paragraph")
    para(doc, (
        "We train with Soft Actor-Critic (SAC; Haarnoja et al., 2018) combined with "
        "Hindsight Experience Replay (HER; Andrychowicz et al., 2017), using the future "
        "relabeling strategy with k=4 virtual transitions per real transition. "
        "The policy and critic are 2-layer MLPs with 256 hidden units. "
        "We use a learning rate of 10⁻³, discount γ=0.98, soft update τ=0.005, "
        "batch size 256, and a replay buffer of 10⁶ transitions. "
        "The entropy temperature α is auto-tuned. "
        "We use Stable-Baselines3 2.8.0 as the SAC backbone, with our own "
        "HERRelabeler implementation for the ablation study."
    ))

    h2(doc, "7.4 Methods — Training Setup Paragraph")
    para(doc, (
        "Each agent is trained for 500,000 environment steps. We run N=8 independent "
        "seeds per experiment. We evaluate the deterministic policy on 100 fresh episodes "
        "every 25,000 training steps using a separate, independently-seeded environment. "
        "We report Interquartile Mean (IQM) and stratified bootstrap 95% confidence "
        "intervals following Agarwal et al. (2021) using the rliable library. "
        "Every run records the exact Git commit, full resolved configuration, and package "
        "versions for reproducibility. Training was performed on an NVIDIA A100-SXM4-80GB GPU."
    ))

    h2(doc, "7.5 Results Table (Fill After Runs)")
    fill(doc, "Fill in after pulling W&B data and computing IQM + 95% CI with rliable.")
    table(doc,
          ["Task", "Steps", "Success Rate (IQM ± 95% CI)", "Target", "Matched?"],
          [
              ["FetchReach", "500k", "[FILL]", "~100%", "[FILL]"],
              ["FetchPush", "500k", "[FILL]", "~80–90%", "[FILL]"],
              ["FetchPickAndPlace", "500k", "[FILL]", "~50–70%", "[FILL]"],
              ["FetchSlide", "500k", "[FILL]", "~30–60%", "[FILL]"],
          ])

    h2(doc, "7.6 Failure Analysis Template")
    para(doc, (
        "Failure analysis. "
    ), bold=True)
    fill(doc, (
        "[FILL — be honest about any seeds that failed to learn, any task where success "
        "rate is below target, and the likely cause. Example: 'FetchSlide showed high "
        "variance across seeds (CI width: [FILL]), consistent with the difficulty of the "
        "task under sparse reward and fixed learning rate. We did not tune hyperparameters "
        "per task.']"
    ))

    # ------------------------------------------------------------------
    h1(doc, "8. Key Citations")
    # ------------------------------------------------------------------

    table(doc,
          ["Work", "Why cite", "BibTeX key"],
          [
              ["Andrychowicz et al. 2017", "HER algorithm, target numbers", "andrychowicz2017her"],
              ["Haarnoja et al. 2018", "SAC algorithm", "haarnoja2018sac"],
              ["Plappert et al. 2018", "GoalEnv interface, Fetch tasks", "plappert2018fetch"],
              ["Raffin et al. 2021", "Stable-Baselines3", "raffin2021sb3"],
              ["Agarwal et al. 2021", "rliable, IQM, statistical reporting", "agarwal2021rliable"],
              ["Todorov et al. 2012", "MuJoCo physics", "todorov2012mujoco"],
              ["Farama Foundation", "Gymnasium-Robotics library", "gymnasium_robotics"],
          ])

    # ------------------------------------------------------------------
    h1(doc, "9. Phase 1 → Paper Checklist")
    # ------------------------------------------------------------------

    checklist(doc, [
        (True,  "Repo scaffold (Phase 0): configs, CI, tests, seeding, logging, provenance"),
        (True,  "HER relabeler implemented and unit-tested (roborto/buffers/her.py)"),
        (True,  "Replay buffer implemented and unit-tested (roborto/buffers/replay_buffer.py)"),
        (True,  "SACHERTrainer with smoke + resume support"),
        (True,  "Smoke run confirmed end-to-end on GPU (make smoke — 2026-05-21)"),
        (True,  "exp001 config set to 500k steps, 8 seeds"),
        (False, "Full 8-seed runs complete and logged to W&B"),
        (False, "Pull W&B data, compute IQM + 95% CI with rliable"),
        (False, "Generate learning-curve figures with scripts/make_figures.py"),
        (False, "Fill in results table (Section 7.5)"),
        (False, "Write failure analysis section"),
        (False, "Append results to docs/lab_notebook.md"),
        (False, "Confirm git SHA of results commit for reproducibility statement"),
        (False, "Extend to FetchPickAndPlace and FetchSlide"),
        (False, "Phase 1 contribution: HER relabeling-strategy sweep (exp002)"),
        (False, "Draft Report 1 in reports/report-01-baselines/main.tex"),
        (False, "Get arXiv endorser lined up (do not leave this to the last week)"),
    ])


def main() -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    build(doc)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
